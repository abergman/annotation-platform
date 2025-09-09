"""
Text Processing Utilities

Functions for processing uploaded files and extracting text content.
"""

import io
import os
from typing import Optional
from fastapi import UploadFile, HTTPException
import docx
import PyPDF2
import pandas as pd


async def process_uploaded_file(file: UploadFile, file_extension: str) -> str:
    """
    Process uploaded file and extract text content.
    
    Args:
        file: Uploaded file object
        file_extension: File extension (e.g., '.txt', '.docx', '.pdf')
    
    Returns:
        str: Extracted text content
    
    Raises:
        HTTPException: If file processing fails
    """
    
    try:
        content = await file.read()
        
        if file_extension == ".txt":
            return process_text_file(content)
        elif file_extension == ".docx":
            return process_docx_file(content)
        elif file_extension == ".pdf":
            return process_pdf_file(content)
        elif file_extension == ".csv":
            return process_csv_file(content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to process file: {str(e)}"
        )
    finally:
        # Reset file pointer
        await file.seek(0)


def process_text_file(content: bytes) -> str:
    """Process plain text file."""
    try:
        # Try UTF-8 first, then fall back to other encodings
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except UnicodeDecodeError:
                return content.decode('utf-8', errors='replace')
    except Exception as e:
        raise ValueError(f"Failed to decode text file: {str(e)}")


def process_docx_file(content: bytes) -> str:
    """Process Microsoft Word document."""
    try:
        doc = docx.Document(io.BytesIO(content))
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise ValueError(f"Failed to process DOCX file: {str(e)}")


def process_pdf_file(content: bytes) -> str:
    """Process PDF document."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        if not text_parts:
            raise ValueError("No text could be extracted from PDF")
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise ValueError(f"Failed to process PDF file: {str(e)}")


def process_csv_file(content: bytes) -> str:
    """Process CSV file by combining all text columns."""
    try:
        # Read CSV into DataFrame
        df = pd.read_csv(io.BytesIO(content))
        
        # Combine all text columns
        text_parts = []
        for index, row in df.iterrows():
            row_text = []
            for col_name, value in row.items():
                if pd.notna(value) and str(value).strip():
                    row_text.append(f"{col_name}: {str(value).strip()}")
            
            if row_text:
                text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            raise ValueError("No text data found in CSV file")
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise ValueError(f"Failed to process CSV file: {str(e)}")


def clean_text(text: str) -> str:
    """
    Clean and normalize text content.
    
    Args:
        text: Raw text content
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if line:  # Skip empty lines
            # Replace multiple spaces with single space
            line = ' '.join(line.split())
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)


def calculate_text_stats(text: str) -> dict:
    """
    Calculate various statistics for text content.
    
    Args:
        text: Text content
    
    Returns:
        dict: Text statistics
    """
    if not text:
        return {
            "word_count": 0,
            "character_count": 0,
            "paragraph_count": 0,
            "sentence_count": 0,
            "average_words_per_sentence": 0
        }
    
    # Basic counts
    word_count = len(text.split())
    character_count = len(text)
    paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
    
    # Rough sentence count (count periods, exclamation marks, question marks)
    sentence_count = text.count('.') + text.count('!') + text.count('?')
    if sentence_count == 0:
        sentence_count = 1  # Avoid division by zero
    
    average_words_per_sentence = round(word_count / sentence_count, 2)
    
    return {
        "word_count": word_count,
        "character_count": character_count,
        "paragraph_count": paragraph_count,
        "sentence_count": sentence_count,
        "average_words_per_sentence": average_words_per_sentence
    }